"""
extractor.py — Extract readable text from uploaded academic articles.
Supports PDF, DOCX, and TXT formats.
"""

import io
import re
from typing import Optional


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """
    Route to the correct extractor based on file extension.
    Returns (extracted_text, info_dict).
    info_dict = {pages, word_count, char_count, success, error}
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _extract_pdf(file_bytes, filename)
    elif ext == "docx":
        return _extract_docx(file_bytes, filename)
    elif ext in ("txt", "md"):
        return _extract_txt(file_bytes, filename)
    else:
        return "", {"pages": 0, "word_count": 0, "char_count": 0,
                    "success": False, "error": f"Unsupported format: .{ext}"}


def _extract_pdf(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = len(doc)
        texts = []
        for page in doc:
            texts.append(page.get_text("text"))
        doc.close()
        raw = "\n".join(texts)
        cleaned = _clean_text(raw)
        if len(cleaned.strip()) < 50:
            return "", {"pages": pages, "word_count": 0, "char_count": 0,
                        "success": False,
                        "error": "No extractable text — this may be a scanned PDF."}
        wc = len(cleaned.split())
        return cleaned, {"pages": pages, "word_count": wc,
                         "char_count": len(cleaned), "success": True, "error": None}
    except Exception as e:
        return "", {"pages": 0, "word_count": 0, "char_count": 0,
                    "success": False, "error": str(e)}


def _extract_docx(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        raw = "\n".join(parts)
        cleaned = _clean_text(raw)
        wc = len(cleaned.split())
        return cleaned, {"pages": max(1, wc // 350), "word_count": wc,
                         "char_count": len(cleaned), "success": True, "error": None}
    except Exception as e:
        return "", {"pages": 0, "word_count": 0, "char_count": 0,
                    "success": False, "error": str(e)}


def _extract_txt(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    try:
        raw = file_bytes.decode("utf-8", errors="replace")
        cleaned = _clean_text(raw)
        wc = len(cleaned.split())
        return cleaned, {"pages": max(1, wc // 350), "word_count": wc,
                         "char_count": len(cleaned), "success": True, "error": None}
    except Exception as e:
        return "", {"pages": 0, "word_count": 0, "char_count": 0,
                    "success": False, "error": str(e)}


def _clean_text(text: str) -> str:
    """Remove noise: headers, footers, page numbers, excessive whitespace."""
    # Remove lines that are just page numbers or short noise
    lines = text.split("\n")
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Skip pure page number lines
        if re.match(r"^\d+$", stripped):
            continue
        # Skip very short lines (likely headers/footers)
        if len(stripped) < 4:
            continue
        clean_lines.append(stripped)
    text = " ".join(clean_lines)
    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)
    # Fix hyphenated line breaks (e.g., "impor-\ntant" → "important")
    text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)
    return text.strip()


def truncate_text(text: str, max_chars: int = 12000) -> tuple[str, bool]:
    """Truncate to max_chars. Returns (text, was_truncated)."""
    if len(text) <= max_chars:
        return text, False
    # Try to cut at sentence boundary
    truncated = text[:max_chars]
    last_period = truncated.rfind(". ")
    if last_period > max_chars * 0.8:
        truncated = truncated[:last_period + 1]
    return truncated, True


def get_first_chars(text: str, n: int = 1000) -> str:
    """Return first n characters for metadata extraction."""
    return text[:n]
